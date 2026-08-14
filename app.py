import io, os, re, json, sqlite3, secrets, urllib.request, urllib.parse, ssl
from datetime import datetime, timezone
from flask import Flask, render_template, request, jsonify, session
from dotenv import load_dotenv
from openai import OpenAI
from pypdf import PdfReader
from docx import Document
load_dotenv(); app=Flask(__name__); app.secret_key=os.getenv('FLASK_SECRET_KEY',secrets.token_hex(32)); app.config['MAX_CONTENT_LENGTH']=8*1024*1024
DB_PATH=os.getenv('DATABASE_PATH','jobhunter.db'); OPENAI_API_KEY=os.getenv('OPENAI_API_KEY',''); OPENAI_MODEL=os.getenv('OPENAI_MODEL','gpt-5-mini'); client=OpenAI(api_key=OPENAI_API_KEY,timeout=20,max_retries=0) if OPENAI_API_KEY else None
def db(): c=sqlite3.connect(DB_PATH); c.row_factory=sqlite3.Row; return c
def init_db():
 c=db(); c.execute('CREATE TABLE IF NOT EXISTS profiles (id TEXT PRIMARY KEY, profile_json TEXT NOT NULL, created_at TEXT NOT NULL)'); c.execute('CREATE TABLE IF NOT EXISTS opportunities (id TEXT PRIMARY KEY, profile_id TEXT NOT NULL, title TEXT, company TEXT, location TEXT, url TEXT, source TEXT, score INTEGER, reason TEXT, created_at TEXT NOT NULL)'); c.execute('CREATE TABLE IF NOT EXISTS feedback (id TEXT PRIMARY KEY, name TEXT, rating INTEGER, message TEXT, created_at TEXT NOT NULL)'); c.commit(); c.close()
init_db()
def clean_text(x): return re.sub(r'\n{3,}','\n\n',re.sub(r'[ \t]+',' ',x.replace('\x00',' '))).strip()
def pdf_ocr(data):
 try:
  import pytesseract; from pdf2image import convert_from_bytes
  return clean_text('\n'.join(pytesseract.image_to_string(p,config='--psm 3') for p in convert_from_bytes(data,dpi=220,fmt='png',thread_count=1)[:12]))
 except Exception as e: app.logger.warning('OCR fallback failed: %s',e); return ''
def extract_cv(file):
 if not file or not file.filename: raise ValueError('Please upload a CV.')
 name=file.filename.lower(); data=file.read()
 if name.endswith('.pdf'):
  text=clean_text('\n'.join(p.extract_text() or '' for p in PdfReader(io.BytesIO(data)).pages))
  if len(re.sub(r'\s','',text))<180 or not re.search(r'(?i)[\w.+-]+\s*@\s*[\w.-]+\s*\.\s*[a-z]{2,}',text) or not re.search(r'(?<!\d)(?:\+|00)?[\d() .-]{8,20}\d(?!\d)',text): text=(text+'\n'+pdf_ocr(data)).strip()
 elif name.endswith('.docx'):
  doc=Document(io.BytesIO(data)); parts=[p.text for p in doc.paragraphs]
  for table in doc.tables:
   for row in table.rows: parts.append(' | '.join(c.text for c in row.cells))
  text=clean_text('\n'.join(parts))
 elif name.endswith('.txt'): text=clean_text(data.decode('utf-8',errors='ignore'))
 else: raise ValueError('Use PDF, DOCX or TXT.')
 if len(text)<40: raise ValueError('We could not read enough text from this CV.')
 return text[:60000]
def ai_json(instructions,prompt):
 if not client: raise RuntimeError('OPENAI_API_KEY is not configured in Railway Variables.')
 return json.loads(client.responses.create(model=OPENAI_MODEL,instructions=instructions,input=prompt,text={'format':{'type':'json_object'}},store=False).output_text)
def detect_contacts(text):
 srcs=[text,re.sub(r'\s*([@.])\s*',r'\1',text)]; email=''
 for src in srcs:
  for x in re.findall(r'(?i)(?<![\w.+-])[A-Z0-9._%+-]+\s*@\s*[A-Z0-9.-]+\s*\.\s*[A-Z]{2,}(?![\w.-])',src):
   x=re.sub(r'\s+','',x).strip('.,;:()[]{}<>')
   if '@' in x and '.' in x.split('@',1)[1]: email=x; break
  if email: break
 lines=text.splitlines(); labeled=[x for x in lines if re.search(r'(?i)\b(phone|mobile|cell|tel|telephone|contact|whatsapp|viber)\b',x)]; sources=labeled+[x for x in lines if x not in labeled]
 for line in sources:
  for cand in re.findall(r'(?<!\d)(?:\+\d{1,3}[\s().-]?)?(?:\d[\s().-]?){8,14}\d(?!\d)',line):
   digits=re.sub(r'\D','',cand)
   if 8<=len(digits)<=15 and len(set(digits))>2 and not re.search(r'(?i)\b(19|20)\d{2}\s*[-–—]\s*(19|20)\d{2}\b',line): return email,re.sub(r'\s+',' ',cand).strip(' -.,;:')
 return email,''
def local_profile(text,email,phone):
 lines=[re.sub(r'\s+',' ',x).strip() for x in text.splitlines() if x.strip()]; first=last=''
 for line in lines[:20]:
  clean=re.sub(r"[^A-Za-zÀ-ž'’\- ]",'',line).strip(); parts=clean.split()
  if 2<=len(parts)<=4 and all(len(p)>1 for p in parts) and not re.search(r'(?i)cv|resume|curriculum|email|phone|linkedin|profile|experience',clean): first,last=parts[0],parts[-1]; break
 exp=[x[:180] for x in lines if re.search(r'(?i)\b(20\d{2}|19\d{2})\b',x) and (re.search(r'(?i)experience|work|employment|present|current',x) or re.search(r'[-–—]',x))]
 years_found=[int(x) for x in re.findall(r'(?<!\d)(?:19|20)(\d{2})(?!\d)',text)]; years=max(0,min(40,datetime.now().year-min(1900+y if y<30 else 2000+y for y in years_found))) if years_found else 0
 titles=[x for x in lines if re.search(r'(?i)\b(manager|analyst|specialist|engineer|developer|consultant|agent|support|assistant|coordinator|administrator|technician|designer|sales|customer service|service desk|help desk|desktop support)\b',x) and len(x)<100]
 return {'first_name':first,'last_name':last,'email':email,'phone':phone,'location':'','job_titles':list(dict.fromkeys(titles))[:10],'years_experience':years,'skills':[],'languages':[],'education':[],'work_experience':exp[:8],'missing_required':[],'profile_ready':False}
def parse_profile(text):
 e,p=detect_contacts(text); x=local_profile(text,e,p); x['missing_required']=[k for k,v in [('email',e),('phone',p),('first_name',x['first_name']),('last_name',x['last_name'])] if not v]+(['work_experience'] if not x['work_experience'] else []); x['profile_ready']=not x['missing_required']; return x
def fetch_json(url):
 with urllib.request.urlopen(urllib.request.Request(url,headers={'User-Agent':'JobHunterAI/1.0','Accept':'application/json'}),timeout=10,context=ssl.create_default_context()) as r:return json.loads(r.read().decode('utf-8',errors='ignore'))
COUNTRY_TERMS={'lithuania':['lithuania','vilnius','kaunas','klaipeda','siauliai','panevezys','šiauliai','panevėžys'],'philippines':['philippines','manila','taguig','cebu','makati','quezon city','pasig','davao'],'germany':['germany','berlin','munich','hamburg','frankfurt','cologne'],'norway':['norway','oslo','bergen','stavanger','trondheim'],'thailand':['thailand','bangkok','chiang mai','phuket','pattaya'],'united kingdom':['united kingdom','uk','london','manchester','birmingham','edinburgh'],'usa':['usa','united states','new york','california','texas','florida']}
COUNTRY_CODE={'lithuania':'lt','germany':'de','norway':'no','thailand':'th','philippines':'ph','united kingdom':'gb','usa':'us'}
def discover_jobs(location,remote,titles):
 jobs=[]; loc=(location or '').strip().lower(); terms=COUNTRY_TERMS.get(loc,[loc] if loc else [])
 try:
  data=fetch_json('https://www.arbeitnow.com/api/job-board-api')
  for j in data.get('data',[]):
   hay=' '.join(str(j.get(k,'')) for k in ['title','description','company_name','location','tags']).lower()
   if loc and not remote and terms and not any(t in hay for t in terms): continue
   if remote and not (j.get('remote') or 'remote' in hay): continue
   jobs.append({'title':j.get('title',''),'company':j.get('company_name',''),'location':j.get('location',''),'url':j.get('url',''),'source':'Arbeitnow'})
 except Exception as e: app.logger.warning('Arbeitnow failed: %s',e)
 if remote:
  try:
   data=fetch_json('https://remotive.com/api/remote-jobs')
   for j in data.get('jobs',[]):
    hay=' '.join(str(j.get(k,'')) for k in ['title','description','company_name','candidate_required_location']).lower(); req=str(j.get('candidate_required_location','')).lower()
    if loc and loc!='worldwide' and terms and not any(t in hay or t in req for t in terms): continue
    jobs.append({'title':j.get('title',''),'company':j.get('company_name',''),'location':j.get('candidate_required_location','Remote'),'url':j.get('url',''),'source':'Remotive'})
  except Exception as e: app.logger.warning('Remotive failed: %s',e)
  try:
   geo=COUNTRY_CODE.get(loc,''); url='https://jobicy.com/api/v2/remote-jobs?count=100'+(('&geo='+geo) if geo else '')
   data=fetch_json(url)
   for j in data.get('jobs',[]): jobs.append({'title':j.get('jobTitle',''),'company':j.get('companyName',''),'location':j.get('jobGeo','Remote'),'url':j.get('url',''),'source':'Jobicy'})
  except Exception as e: app.logger.warning('Jobicy failed: %s',e)
 app_id=os.getenv('ADZUNA_APP_ID',''); app_key=os.getenv('ADZUNA_APP_KEY',''); code=COUNTRY_CODE.get(loc,'')
 if app_id and app_key and code:
  try:
   for page in range(1,4):
    url=f'https://api.adzuna.com/v1/api/jobs/{code}/search/{page}?app_id={urllib.parse.quote(app_id)}&app_key={urllib.parse.quote(app_key)}&results_per_page=50&content-type=application/json'; data=fetch_json(url)
    for j in data.get('results',[]): jobs.append({'title':j.get('title',''),'company':(j.get('company') or {}).get('display_name',''),'location':(j.get('location') or {}).get('display_name',''),'url':j.get('redirect_url',''),'source':'Adzuna'})
    if len(data.get('results',[]))<50: break
  except Exception as e: app.logger.warning('Adzuna failed: %s',e)
 seen=set(); out=[]
 for j in jobs:
  title=re.sub(r'\s+',' ',j['title']).strip(); company=re.sub(r'\s+',' ',j['company']).strip(); url=j.get('url','').strip(); key=(re.sub(r'\W','',title.lower()),re.sub(r'\W','',company.lower()))+((url.split('?')[0],) if url else ())
  if title and key not in seen: seen.add(key); j['title']=title; j['company']=company; out.append(j)
 return out[:500]
def _tokens(s): return set(re.findall(r'[a-z0-9+#.]{2,}',str(s).lower()))
def _deterministic_score(profile,job):
 title=_tokens(job.get('title','')); text=_tokens(' '.join(str(job.get(k,'')) for k in ['title','company','location']))
 profile_terms=set()
 for v in profile.get('job_titles',[]): profile_terms |= _tokens(v)
 for v in profile.get('skills',[]): profile_terms |= _tokens(v)
 for v in profile.get('languages',[]): profile_terms |= _tokens(v)
 for v in profile.get('work_experience',[]): profile_terms |= _tokens(v)
 if not profile_terms: return 25,'Job found; CV details are limited, so review the role manually.'
 exact=len(title & profile_terms); related=0
 groups=[{'support','service','help','desk','technical','it','desktop','customer'},{'analyst','specialist','consultant','administrator','technician'},{'manager','lead','supervisor'},{'developer','engineer','software','data','cloud','security'}]
 for g in groups:
  if title & g and profile_terms & g: related += 1
 score=min(95,25 + exact*9 + related*12)
 if exact==0 and related==0: score=12
 reason='Strong title/skill overlap.' if score>=70 else ('Some relevant experience overlap.' if score>=45 else 'Potentially relevant; review requirements.')
 return score,reason
def match_jobs(profile,jobs):
 if not jobs:return []
 # Always produce a useful non-zero baseline locally. AI is an enhancement, never the source of truth.
 base=[]
 for j in jobs:
  s,r=_deterministic_score(profile,j); base.append({'job':j,'score':s,'reason':r})
 compact=[{'index':i,'title':x['job']['title'],'company':x['job']['company'],'location':x['job']['location'],'source':x['job']['source']} for i,x in enumerate(base[:120])]
 if client:
  try:
   result=ai_json('You are a recruitment matching engine. Score real relevance from the candidate profile. Never invent qualifications. A zero score should be used only when the job is clearly unrelated. Return every index provided. Keep reasons under 16 words.',f'Candidate profile: {json.dumps(profile)} Jobs: {json.dumps(compact)} Return JSON exactly: {{"matches":[{{"index":0,"score":50,"reason":"short truthful reason"}}]}}. Score 0-100 based on title, skills, experience and languages. Do not score every job 0.')
   for x in result.get('matches',[]):
    try:
     i=int(x.get('index')); s=int(x.get('score'))
     if 0<=i<len(base): base[i]['score']=max(1,min(100,s)); base[i]['reason']=str(x.get('reason') or base[i]['reason'])[:240]
    except (ValueError,TypeError): pass
  except Exception as e: app.logger.warning('AI matching unavailable; using deterministic scores: %s',e)
 return base
def application_draft(profile,job): return ai_json('You write concise professional job applications using only supported facts.',f'Create a short truthful application email. Candidate: {json.dumps(profile)} Job: {json.dumps(job)} Return JSON: {{"subject":"","body":""}}')
@app.get('/')
def home(): return render_template('index.html')
@app.after_request
def inject(response):
 if response.content_type and 'text/html' in response.content_type and response.status_code==200:
  b=response.get_data(as_text=True); tags='<script src="/static/enhancements.js"></script><script src="/static/api-fix.js"></script><script src="/static/loading.js"></script>'
  if '/static/api-fix.js' not in b:b=b.replace('</body>',tags+'</body>')
  response.set_data(b)
 return response
@app.get('/health')
def health(): return jsonify(status='ok',service='jobhunter-ai')
@app.post('/api/profile')
def profile():
 try:
  p=parse_profile(extract_cv(request.files.get('cv'))); pid=secrets.token_urlsafe(16); c=db(); c.execute('INSERT INTO profiles VALUES (?,?,?)',(pid,json.dumps(p),datetime.now(timezone.utc).isoformat())); c.commit(); c.close(); session['profile_id']=pid; return jsonify({'profile_id':pid,**p})
 except Exception as e: app.logger.exception('profile failed'); return jsonify(error=str(e)),400
@app.post('/api/search')
def search():
 try:
  pid=session.get('profile_id'); c=db(); row=c.execute('SELECT profile_json FROM profiles WHERE id=?',(pid,)).fetchone() if pid else None
  if not row: c.close(); return jsonify(error='Upload your CV first.'),400
  p=json.loads(row['profile_json']); location=request.form.get('location','').strip(); remote=request.form.get('remote')=='true'; jobs=discover_jobs(location,remote,p.get('job_titles',[]))
  if not jobs: c.close(); return jsonify(opportunities=[],count=0,message='No public listings were returned. Try enabling Remote or another country.')
  matches=match_jobs(p,jobs); out=[]
  for item in matches:
   j=item['job']; oid=secrets.token_urlsafe(16); c.execute('INSERT INTO opportunities VALUES (?,?,?,?,?,?,?,?,?,?)',(oid,pid,j['title'],j['company'],j['location'],j['url'],j['source'],item['score'],item['reason'],datetime.now(timezone.utc).isoformat())); out.append({'id':oid,**j,'score':item['score'],'reason':item['reason']})
  c.commit(); c.close(); out.sort(key=lambda x:x['score'],reverse=True); return jsonify(opportunities=out[:50],count=len(out),sources=sorted(set(x['source'] for x in out)))
 except Exception as e: app.logger.exception('search failed'); return jsonify(error=str(e)),500
@app.post('/api/draft')
def draft():
 try:
  pid=session.get('profile_id'); oid=request.form.get('opportunity_id'); c=db(); r=c.execute('SELECT * FROM opportunities WHERE id=? AND profile_id=?',(oid,pid)).fetchone(); p=c.execute('SELECT profile_json FROM profiles WHERE id=?',(pid,)).fetchone(); c.close()
  if not r or not p:return jsonify(error='Opportunity not found.'),404
  return jsonify(**application_draft(json.loads(p['profile_json']),dict(r)))
 except Exception as e:return jsonify(error=str(e)),400
@app.post('/api/chat')
def chat():
 try:
  payload=request.get_json(silent=True) or {}; msg=(payload.get('message') or request.form.get('message','')).strip()
  if not msg:return jsonify(error='Please type a question.'),400
  pid=session.get('profile_id'); c=db(); row=c.execute('SELECT profile_json FROM profiles WHERE id=?',(pid,)).fetchone() if pid else None; c.close(); prof=json.loads(row['profile_json']) if row else {}
  a=ai_json('You are JobHunter AI. Be concise and direct: 1-4 short sentences or up to 4 bullets, under 60 words.',f'Profile: {json.dumps(prof)} Question: {msg} Return JSON: {{"answer":""}}'); return jsonify(answer=str(a.get('answer',''))[:700])
 except Exception as e:return jsonify(error=str(e)),500
@app.post('/api/feedback')
def feedback():
 try:
  name=request.form.get('name','')[:80]; msg=request.form.get('message','')[:1000]; rating=max(1,min(5,int(request.form.get('rating','5'))))
  if not msg:return jsonify(error='Please write a short message.'),400
  c=db(); c.execute('INSERT INTO feedback VALUES (?,?,?,?,?)',(secrets.token_urlsafe(12),name,rating,msg,datetime.now(timezone.utc).isoformat())); c.commit(); c.close(); return jsonify(ok=True)
 except Exception as e:return jsonify(error=str(e)),400
@app.errorhandler(404)
def not_found(e):
 if request.path.startswith('/api/'): return jsonify(error='Not found'),404
 return render_template('index.html'),404
@app.errorhandler(413)
def too_large(e): return jsonify(error='CV is too large. Maximum size is 8 MB.'),413
if __name__=='__main__': app.run(host='0.0.0.0',port=int(os.getenv('PORT','5000')))