"""Production WSGI wrapper for robust CV extraction.

The original Flask app remains the product application. This wrapper adds a
seven-layer deterministic extraction pipeline before the existing /api/profile
handler runs, without putting OpenAI on the upload critical path.
"""
import io
import re
from pathlib import Path

import app as application
from pypdf import PdfReader
from docx import Document


def _clean(s: str) -> str:
    s = (s or '').replace('\x00', ' ')
    s = s.replace('\u00ad', '')
    s = re.sub(r'[ \t]+', ' ', s)
    s = re.sub(r'\n{3,}', '\n\n', s)
    return s.strip()


def _ocr_pdf(data: bytes) -> str:
    try:
        import pytesseract
        from pdf2image import convert_from_bytes
        pages = convert_from_bytes(data, dpi=260, fmt='png', thread_count=1, first_page=1, last_page=12)
        chunks = []
        for page in pages:
            chunks.append(pytesseract.image_to_string(page, config='--psm 3'))
        return _clean('\n'.join(chunks))
    except Exception as exc:
        application.app.logger.warning('OCR fallback failed: %s', exc)
        return ''


def _docx_text(data: bytes) -> str:
    doc = Document(io.BytesIO(data))
    chunks = [p.text for p in doc.paragraphs]
    for section in doc.sections:
        chunks.extend(p.text for p in section.header.paragraphs)
        chunks.extend(p.text for p in section.footer.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            chunks.append(' | '.join(cell.text for cell in row.cells))
    return _clean('\n'.join(chunks))


def robust_extract(file):
    if not file or not file.filename:
        raise ValueError('Please upload a CV.')
    filename = file.filename.lower()
    data = file.read()
    if len(data) > 8 * 1024 * 1024:
        raise ValueError('CV is too large. Maximum size is 8 MB.')

    if filename.endswith('.pdf'):
        # Layer 1: native PDF text.
        native_pages = []
        for page in PdfReader(io.BytesIO(data)).pages:
            native_pages.append(page.extract_text() or '')
        native = _clean('\n'.join(native_pages))

        # Layers 2-4: inspect native extraction, then OCR only when it looks
        # incomplete. OCR is also useful for highly visual/two-column CVs.
        combined = native
        preliminary = _detect_contacts(combined)
        if len(re.sub(r'\s', '', native)) < 180 or not preliminary['email'] or not preliminary['phone']:
            ocr = _ocr_pdf(data)
            if ocr:
                combined = _clean(native + '\n' + ocr)
    elif filename.endswith('.docx'):
        # Layer 2: paragraphs + tables + headers/footers.
        combined = _docx_text(data)
    elif filename.endswith('.txt'):
        combined = _clean(data.decode('utf-8', errors='ignore'))
    else:
        raise ValueError('Use PDF, DOCX or TXT.')

    if len(combined) < 40:
        raise ValueError('We could not read enough text from this CV. Please upload a clearer PDF or DOCX.')
    return combined[:60000]


def _email_candidates(text: str):
    # Layer 5: tolerate spaces and common OCR substitutions around @ and dot.
    srcs = [text]
    normalized = re.sub(r'(?i)\[\s*(at|dot)\s*\]', lambda m: '@' if m.group(1).lower() == 'at' else '.', text)
    normalized = re.sub(r'(?i)\s+(?:at)\s+', '@', normalized)
    normalized = re.sub(r'(?i)\s+(?:dot)\s+', '.', normalized)
    normalized = re.sub(r'\s*([@.])\s*', r'\1', normalized)
    srcs.append(normalized)
    pattern = re.compile(r'(?i)(?<![\w.+-])[A-Z0-9._%+-]+\s*@\s*[A-Z0-9.-]+\s*\.\s*[A-Z]{2,}(?![\w.-])')
    found = []
    for src in srcs:
        found.extend(pattern.findall(src))
    return found


def _valid_email(value: str) -> bool:
    value = re.sub(r'\s+', '', value).strip('.,;:()[]{}<>')
    if not (5 <= len(value) <= 254 and value.count('@') == 1):
        return False
    local, domain = value.rsplit('@', 1)
    return bool(local and '.' in domain and not domain.startswith('.') and not domain.endswith('.') and '..' not in value)


def _phone_candidates(text: str):
    # Layer 5: broad international phone candidates. Validation below rejects
    # years, dates, ZIP/postal codes and low-information numeric strings.
    labeled = [x for x in text.splitlines() if re.search(r'(?i)\b(phone|mobile|cell|tel|telephone|contact|whatsapp|viber|mob)\b', x)]
    other = [x for x in text.splitlines() if x not in labeled]
    sources = labeled + other
    patterns = [
        r'(?<!\d)(?:\+|00)?\s*\d{1,3}(?:[\s().-]*\d){7,14}(?!\d)',
        r'(?<!\d)\(?\d{2,4}\)?[\s.-]\d{2,4}[\s.-]\d{2,5}(?!\d)',
    ]
    for line in sources:
        for pattern in patterns:
            for candidate in re.findall(pattern, line):
                digits = re.sub(r'\D', '', candidate)
                if not 8 <= len(digits) <= 15:
                    continue
                if len(set(digits)) <= 2:
                    continue
                # Avoid date/year ranges and obvious dates.
                if re.search(r'(?i)\b(?:19|20)\d{2}\s*[-–—]\s*(?:19|20)\d{2}\b', line):
                    continue
                if re.fullmatch(r'(?:19|20)\d{6}', digits):
                    continue
                return re.sub(r'\s+', ' ', candidate).strip(' -.,;:')
    return ''


def _detect_contacts(text: str):
    # Layer 6: deterministic validation and normalization.
    email = ''
    for candidate in _email_candidates(text):
        candidate = re.sub(r'\s+', '', candidate).strip('.,;:()[]{}<>')
        if _valid_email(candidate):
            email = candidate
            break
    phone = _phone_candidates(text)

    # Layer 7: confidence scoring. AI is deliberately NOT called here; contact
    # extraction must never hang the upload request.
    email_conf = 0.0
    phone_conf = 0.0
    if email:
        email_conf = 0.99 if '@' in email and '.' in email.rsplit('@', 1)[1] else 0.85
        if re.search(r'(?i)\b(email|e-mail|mail)\b', text[:4000]):
            email_conf = min(1.0, email_conf + 0.005)
    if phone:
        digits = re.sub(r'\D', '', phone)
        phone_conf = 0.96 if phone.lstrip().startswith('+') or phone.lstrip().startswith('00') else 0.91
        if re.search(r'(?i)\b(phone|mobile|cell|tel|telephone|contact|whatsapp|viber|mob)\b', text):
            phone_conf = min(1.0, phone_conf + 0.04)
    return {'email': email, 'phone': phone, 'email_confidence': round(email_conf, 3), 'phone_confidence': round(phone_conf, 3)}


def robust_detect_contacts(text):
    result = _detect_contacts(text)
    return result['email'], result['phone']


def robust_parse_profile(text):
    contacts = _detect_contacts(text)
    # Reuse the application's conservative local profile builder for the rest.
    profile = application.local_profile(text, contacts['email'], contacts['phone'])
    profile['email_confidence'] = contacts['email_confidence']
    profile['phone_confidence'] = contacts['phone_confidence']

    missing = []
    if not profile.get('email'):
        missing.append('email')
    if not profile.get('phone'):
        missing.append('phone')
    if not profile.get('first_name'):
        missing.append('first_name')
    if not profile.get('last_name'):
        missing.append('last_name')
    if not profile.get('work_experience'):
        missing.append('work_experience')
    profile['missing_required'] = missing
    profile['profile_ready'] = not bool(missing)
    return profile


# Monkey-patch only the two upload-stage functions. Search, drafting and chat
# continue using the original application implementation.
application.extract_cv = robust_extract
application.detect_contacts = robust_detect_contacts
application.parse_profile = robust_parse_profile

app = application.app
